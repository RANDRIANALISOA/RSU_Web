"""rapport_word.py — Export WORD (.docx) du rapport de mission RÉDIGÉ PAR IA.

Convertit le Markdown produit par `rapport_mission.synthese_ia_iter` (déjà
ANONYMISÉ : aucun nom de personne) en un document Word moderne et illustré :

  - page de garde (bannière INSTAT, titre, période, périmètre, bande aux couleurs
    du drapeau malgache) ;
  - bandeau d'indicateurs clés (KPI) sous forme de cartes colorées ;
  - deux GRAPHIQUES matplotlib (activité par district, activité par jour) ;
  - le corps rédigé par l'IA (titres, listes, tableaux, gras/italique) avec des
    styles colorés (titres soulignés, couleurs de la charte).

Dépendances serveur : python-docx, matplotlib (backend Agg), Pillow — installées
dans le venv. AUCUN accès réseau ; tout est produit en mémoire (renvoie des octets).

Le module est TOLÉRANT : si une brique optionnelle échoue (bannière absente,
graphique impossible faute de données), on la saute sans casser le document.
"""
import datetime
import io
import os
import re

import matplotlib
matplotlib.use("Agg")               # backend non-interactif (serveur sans écran)
import matplotlib.pyplot as plt
from matplotlib.ticker import MaxNLocator

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import config

# --- Charte graphique (mêmes teintes que le rapport HTML) -------------------
_PRIMAIRE = RGBColor(0x12, 0x32, 0x5C)      # bleu nuit (titres)
_ACCENT = RGBColor(0x15, 0x58, 0xC9)        # bleu vif (chiffres, accents)
_GRIS = RGBColor(0x4B, 0x55, 0x63)          # texte courant
_GRIS_CLAIR = RGBColor(0x64, 0x74, 0x8B)    # légendes
_FLAG_ROUGE = RGBColor(0xFC, 0x3D, 0x32)    # drapeau Madagascar
_FLAG_VERT = RGBColor(0x00, 0x7E, 0x3A)
_HEX_PRIMAIRE = "12325C"
_HEX_ACCENT = "1558C9"
_HEX_CARTE = "F1F5FB"                        # fond clair des cartes KPI
_MPL_PRIMAIRE = "#12325c"
_MPL_ACCENT = "#1558c9"
_MPL_ROUGE = "#e0413a"
_POLICE = "Calibri"


# ---------------------------------------------------------------------------
# Petites aides XML (python-docx n'expose pas tout en direct)
# ---------------------------------------------------------------------------
def _shade(element, hex_couleur):
    """Applique un fond (w:shd) à une cellule ou un paragraphe (via son <…>Pr)."""
    pr = element.get_or_add_tcPr() if element.tag.endswith("}tc") else element
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_couleur)
    pr.append(shd)


def _bordure_bas(paragraphe, hex_couleur=_HEX_PRIMAIRE, taille="8"):
    """Trait horizontal sous un paragraphe (titres de section soulignés)."""
    p_pr = paragraphe._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), taille)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), hex_couleur)
    borders.append(bottom)
    p_pr.append(borders)


def _cell_shade(cell, hex_couleur):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_couleur)
    tc_pr.append(shd)


def _numeros_de_page(doc):
    """Ajoute « Page X » au pied de page (champ Word natif)."""
    section = doc.sections[0]
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("RSU 2026 — Rapport de mission · page ").font.size = Pt(8)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), "16"); rpr.append(sz)
    run.append(rpr)
    fld.append(run)
    p._p.append(fld)


# ---------------------------------------------------------------------------
# Graphiques matplotlib -> PNG en mémoire
# ---------------------------------------------------------------------------
def _fig_vers_png(fig):
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight",
                facecolor="white")
    plt.close(fig)
    buf.seek(0)
    return buf


def _style_axes(ax):
    for cote in ("top", "right"):
        ax.spines[cote].set_visible(False)
    for cote in ("left", "bottom"):
        ax.spines[cote].set_color("#c7d0dc")
    ax.tick_params(colors="#4b5563", labelsize=9)
    ax.grid(axis="y", color="#e6ebf2", linewidth=0.8)
    ax.set_axisbelow(True)


def _graphe_par_district(rapport):
    """Barres horizontales : nombre d'entrées de journal par district."""
    ds = [d for d in rapport.get("districts", []) if d.get("n_entrees")]
    if not ds:
        return None
    ds = sorted(ds, key=lambda d: d["n_entrees"])
    noms = [d["nom"] for d in ds]
    vals = [d["n_entrees"] for d in ds]
    fig, ax = plt.subplots(figsize=(6.4, max(1.6, 0.5 * len(ds) + 0.8)))
    barres = ax.barh(noms, vals, color=_MPL_ACCENT, height=0.62)
    _style_axes(ax)
    ax.grid(axis="y", visible=False)
    ax.grid(axis="x", color="#e6ebf2", linewidth=0.8)
    ax.xaxis.set_major_locator(MaxNLocator(integer=True))
    for b, v in zip(barres, vals):
        ax.text(b.get_width() + max(vals) * 0.01, b.get_y() + b.get_height() / 2,
                str(v), va="center", ha="left", fontsize=9,
                color=_MPL_PRIMAIRE, fontweight="bold")
    ax.set_xlabel("Entrées de journal", fontsize=9, color="#4b5563")
    fig.tight_layout()
    return _fig_vers_png(fig)


def _compter_par_jour(rapport):
    """Agrège le nombre d'entrées par date (toutes fonctions/districts confondus)."""
    par_jour = {}
    for d in rapport.get("districts", []):
        for f in d.get("fonctions", []):
            for p in f.get("personnes", []):
                for e in p.get("entrees", []):
                    j = e.get("date_jour")
                    if j:
                        par_jour[j] = par_jour.get(j, 0) + 1
    return par_jour


def _graphe_par_jour(rapport):
    """Barres verticales : activité (entrées) jour par jour."""
    par_jour = _compter_par_jour(rapport)
    if not par_jour:
        return None
    jours = sorted(par_jour)
    vals = [par_jour[j] for j in jours]

    def _court(iso):
        try:
            return datetime.date.fromisoformat(iso).strftime("%d/%m")
        except Exception:
            return iso

    etiquettes = [_court(j) for j in jours]
    fig, ax = plt.subplots(figsize=(6.4, 2.6))
    ax.bar(range(len(jours)), vals, color=_MPL_PRIMAIRE, width=0.66)
    _style_axes(ax)
    ax.set_xticks(range(len(jours)))
    ax.set_xticklabels(etiquettes, rotation=45, ha="right", fontsize=8)
    ax.yaxis.set_major_locator(MaxNLocator(integer=True))
    ax.set_ylabel("Entrées", fontsize=9, color="#4b5563")
    fig.tight_layout()
    return _fig_vers_png(fig)


def _banniere_png():
    """Charge la bannière INSTAT (images/images.jfif) et la renvoie en PNG.
    (python-docx embarque plus sûrement du PNG ; .jfif est un JPEG.)"""
    chemin = os.path.join(config.IMAGES_DIR, "images.jfif")
    if not os.path.exists(chemin):
        return None
    try:
        from PIL import Image
        img = Image.open(chemin).convert("RGB")
        buf = io.BytesIO()
        img.save(buf, format="PNG")
        buf.seek(0)
        return buf
    except Exception:
        return None


# ---------------------------------------------------------------------------
# Markdown -> éléments Word (titres, listes, tableaux, gras/italique)
# ---------------------------------------------------------------------------
_RE_INLINE = re.compile(r"(\*\*.+?\*\*|\*(?!\s).+?\*)")


def _ajouter_runs(paragraphe, texte, base_couleur=_GRIS):
    """Écrit `texte` dans un paragraphe en gérant **gras** et *italique*."""
    for bout in _RE_INLINE.split(texte):
        if not bout:
            continue
        if bout.startswith("**") and bout.endswith("**"):
            r = paragraphe.add_run(bout[2:-2]); r.bold = True
        elif bout.startswith("*") and bout.endswith("*"):
            r = paragraphe.add_run(bout[1:-1]); r.italic = True
        else:
            r = paragraphe.add_run(bout)
        r.font.color.rgb = base_couleur


def _titre(doc, texte, niveau):
    """Titre coloré (soulignement pour les niveaux 1-2)."""
    p = doc.add_paragraph()
    p.space_before = Pt(10)
    run = p.add_run(texte)
    run.bold = True
    run.font.name = _POLICE
    tailles = {1: 16, 2: 13, 3: 11.5, 4: 11}
    run.font.size = Pt(tailles.get(niveau, 11))
    run.font.color.rgb = _PRIMAIRE if niveau <= 2 else _ACCENT
    if niveau <= 2:
        _bordure_bas(p, _HEX_PRIMAIRE if niveau == 1 else _HEX_ACCENT,
                     "10" if niveau == 1 else "6")
    return p


def _tableau_markdown(doc, lignes):
    cellules = [[c.strip() for c in l.strip().strip("|").split("|")] for l in lignes]
    if not cellules:
        return
    entete = cellules[0]
    corps = cellules[1:]
    if len(cellules) > 1 and all(set(c) <= set("-: ") and c for c in cellules[1]):
        corps = cellules[2:]
    tab = doc.add_table(rows=1, cols=len(entete))
    tab.alignment = WD_TABLE_ALIGNMENT.CENTER
    tab.style = "Table Grid"
    for i, txt in enumerate(entete):
        cel = tab.rows[0].cells[i]
        cel.text = ""
        _cell_shade(cel, _HEX_PRIMAIRE)
        r = cel.paragraphs[0].add_run(txt)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(9.5)
    for lig in corps:
        cells = tab.add_row().cells
        for i in range(len(entete)):
            val = lig[i] if i < len(lig) else ""
            cells[i].text = ""
            _ajouter_runs(cells[i].paragraphs[0], val)
            for run in cells[i].paragraphs[0].runs:
                run.font.size = Pt(9.5)


def _ajouter_markdown(doc, md):
    lignes = (md or "").replace("\r\n", "\n").split("\n")
    i, n = 0, len(lignes)
    while i < n:
        s = lignes[i].strip()
        # Tableau
        if s.startswith("|") and s.count("|") >= 2:
            tbl = []
            while i < n and lignes[i].strip().startswith("|"):
                tbl.append(lignes[i].strip())
                i += 1
            _tableau_markdown(doc, tbl)
            continue
        # Titre #..######
        m = re.match(r"^(#{1,6})\s+(.*)$", s)
        if m:
            _titre(doc, m.group(2).replace("**", "").strip(), len(m.group(1)))
            i += 1
            continue
        # Filet horizontal
        if re.match(r"^(-{3,}|\*{3,}|_{3,})$", s):
            i += 1
            continue
        # Listes (ordonnée / à puces)
        ol = re.match(r"^\d+[.)]\s+(.*)$", s)
        ul = re.match(r"^[-*]\s+(.*)$", s)
        if ol or ul:
            style = "List Number" if ol else "List Bullet"
            while i < n:
                ss = lignes[i].strip()
                mm = (re.match(r"^\d+[.)]\s+(.*)$", ss) if ol
                      else re.match(r"^[-*]\s+(.*)$", ss))
                if not mm:
                    break
                p = doc.add_paragraph(style=style)
                _ajouter_runs(p, mm.group(1))
                i += 1
            continue
        # Ligne vide
        if s == "":
            i += 1
            continue
        # Paragraphe (regroupe les lignes contiguës)
        buf = []
        while i < n and lignes[i].strip() and not lignes[i].strip().startswith(
                ("#", "|", "- ", "* ")) and not re.match(r"^\d+[.)]\s", lignes[i].strip()):
            buf.append(lignes[i].strip())
            i += 1
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        _ajouter_runs(p, " ".join(buf))
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


# ---------------------------------------------------------------------------
# Page de garde + KPI + graphiques
# ---------------------------------------------------------------------------
def _bande_drapeau(doc):
    """Fine bande blanc / rouge / vert (drapeau de Madagascar) en accent."""
    tab = doc.add_table(rows=1, cols=3)
    tab.alignment = WD_TABLE_ALIGNMENT.CENTER
    largeurs = [Inches(2.0), Inches(2.2), Inches(2.2)]
    couleurs = [None, "FC3D32", "007E3A"]
    for cel, larg, coul in zip(tab.rows[0].cells, largeurs, couleurs):
        cel.width = larg
        cel.text = ""
        cel.paragraphs[0].add_run(" ").font.size = Pt(4)
        if coul:
            _cell_shade(cel, coul)
    # Pas de bordures visibles sur la bande.
    for cel in tab.rows[0].cells:
        tc_pr = cel._tc.get_or_add_tcPr()
        borders = OxmlElement("w:tcBorders")
        for bord in ("top", "left", "bottom", "right"):
            b = OxmlElement(f"w:{bord}")
            b.set(qn("w:val"), "nil")
            borders.append(b)
        tc_pr.append(borders)


def _page_de_garde(doc, rapport, perim_label):
    banniere = _banniere_png()
    if banniere:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        try:
            p.add_run().add_picture(banniere, width=Inches(6.2))
        except Exception:
            pass

    pays = doc.add_paragraph()
    pays.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pays.space_before = Pt(18)
    r = pays.add_run("RÉPUBLIQUE DE MADAGASCAR — INSTAT")
    r.font.size = Pt(11)
    r.font.color.rgb = _GRIS_CLAIR
    r.bold = True

    titre = doc.add_paragraph()
    titre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = titre.add_run("Rapport de mission")
    r.font.size = Pt(30); r.bold = True; r.font.color.rgb = _PRIMAIRE
    sous = doc.add_paragraph()
    sous.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sous.add_run("Registre Social Unique (RSU) 2026")
    r.font.size = Pt(15); r.font.color.rgb = _ACCENT; r.bold = True

    doc.add_paragraph().add_run().font.size = Pt(4)
    _bande_drapeau(doc)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.space_before = Pt(20)
    r = meta.add_run(f"Période du {rapport['debut']} au {rapport['fin']}")
    r.font.size = Pt(12); r.bold = True; r.font.color.rgb = _GRIS
    meta.add_run("\n")
    r = meta.add_run(f"Périmètre : {perim_label}")
    r.font.size = Pt(11); r.font.color.rgb = _GRIS
    meta.add_run("\n")
    r = meta.add_run(f"Document établi le {rapport['genere_le']}")
    r.font.size = Pt(10); r.font.color.rgb = _GRIS_CLAIR; r.italic = True

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.space_before = Pt(28)
    r = note.add_run("Rapport rédigé automatiquement par intelligence artificielle "
                     "à partir des journaux de bord (sans les noms des personnes). "
                     "À relire et valider avant diffusion.")
    r.font.size = Pt(9); r.italic = True; r.font.color.rgb = _GRIS_CLAIR
    doc.add_page_break()


def _cartes_kpi(doc, rapport):
    s = rapport.get("stats", {})
    items = [
        (str(s.get("n_entrees", 0)), "Entrées"),
        (str(s.get("n_redacteurs", 0)), "Rédacteurs"),
        (f'{s.get("n_jours_couverts", 0)}/{s.get("n_jours_periode", 0)}', "Jours couverts"),
        (str(s.get("n_districts", 0)), "Districts"),
        (str(s.get("n_fonctions", 0)), "Fonctions"),
    ]
    tab = doc.add_table(rows=2, cols=len(items))
    tab.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (val, lib) in enumerate(items):
        c_haut = tab.rows[0].cells[i]
        c_bas = tab.rows[1].cells[i]
        _cell_shade(c_haut, _HEX_CARTE)
        _cell_shade(c_bas, _HEX_CARTE)
        c_haut.text = ""; c_bas.text = ""
        ph = c_haut.paragraphs[0]; ph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rr = ph.add_run(val); rr.bold = True; rr.font.size = Pt(20)
        rr.font.color.rgb = _PRIMAIRE
        pb = c_bas.paragraphs[0]; pb.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rr = pb.add_run(lib.upper()); rr.font.size = Pt(8)
        rr.font.color.rgb = _GRIS_CLAIR; rr.bold = True


def _bloc_graphiques(doc, rapport):
    g_dist = _graphe_par_district(rapport)
    g_jour = _graphe_par_jour(rapport)
    if not (g_dist or g_jour):
        return
    _titre(doc, "Aperçu chiffré de l'activité", 2)
    if g_dist:
        lg = doc.add_paragraph()
        lg.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = lg.add_run("Entrées de journal par district")
        r.bold = True; r.font.size = Pt(10); r.font.color.rgb = _ACCENT
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        try:
            p.add_run().add_picture(g_dist, width=Inches(6.0))
        except Exception:
            pass
    if g_jour:
        lg = doc.add_paragraph()
        lg.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lg.space_before = Pt(6)
        r = lg.add_run("Activité jour par jour (nombre d'entrées)")
        r.bold = True; r.font.size = Pt(10); r.font.color.rgb = _ACCENT
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        try:
            p.add_run().add_picture(g_jour, width=Inches(6.0))
        except Exception:
            pass


# ---------------------------------------------------------------------------
# Point d'entrée
# ---------------------------------------------------------------------------
def construire_docx(markdown, rapport, perim_label) -> bytes:
    """Assemble le document Word et renvoie ses octets (.docx)."""
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = _POLICE
    normal.font.size = Pt(11)
    normal.font.color.rgb = _GRIS
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    _page_de_garde(doc, rapport, perim_label)
    _cartes_kpi(doc, rapport)
    _bloc_graphiques(doc, rapport)
    doc.add_paragraph().add_run().font.size = Pt(4)
    _ajouter_markdown(doc, markdown)
    try:
        _numeros_de_page(doc)
    except Exception:
        pass

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def nom_fichier(rapport) -> str:
    """Nom de fichier proposé au téléchargement."""
    d = (rapport.get("debut") or "").replace("-", "")
    f = (rapport.get("fin") or "").replace("-", "")
    return f"Rapport_mission_RSU2026_{d}_{f}.docx"
